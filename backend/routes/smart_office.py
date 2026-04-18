from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from motor.motor_asyncio import AsyncIOMotorDatabase
from datetime import datetime
import secrets, io, json
from typing import Optional
from database import get_db
from routes.auth import get_current_user
from services.ai import (
    translate_text, analyze_document, excel_nlq,
    generate_document_report, REPORT_PROMPTS,
    analyze_excel_data, detect_data_errors, generate_formula,
    smart_files_chat
)
from services.transcription import upload_audio, transcribe_url

router = APIRouter()

class TranslateRequest(BaseModel):
    text: str
    target_language: str
    source_language: str = "auto"

class ExcelQueryRequest(BaseModel):
    file_id: str
    question: str

class DocumentQueryRequest(BaseModel):
    file_id: str
    question: str = ""

class DocumentReportRequest(BaseModel):
    file_id: str
    report_type: str

class FormulaRequest(BaseModel):
    description: str
    context: str = ""

class SmartFilesChatRequest(BaseModel):
    message: str

class CreateFolderRequest(BaseModel):
    name: str
    parent_id: Optional[str] = None

# ── Translate ────────────────────────────────────────────────────────────────

@router.post("/translate")
async def translate(req: TranslateRequest, current_user: dict = Depends(get_current_user)):
    try:
        result = await translate_text(req.text, req.target_language)
        return {"translated": result, "source": req.text, "target_language": req.target_language}
    except Exception as e:
        raise HTTPException(502, f"Translation service error: {str(e)}")

# ── Documents ─────────────────────────────────────────────────────────────────

@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    try:
        content = await file.read()
        text_content = ""

        if file.filename and file.filename.endswith(".pdf"):
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                text_content = "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception:
                text_content = content.decode("utf-8", errors="ignore")
        elif file.filename and (file.filename.endswith(".docx") or file.filename.endswith(".doc")):
            try:
                from docx import Document
                doc = Document(io.BytesIO(content))
                text_content = "\n".join(para.text for para in doc.paragraphs)
            except Exception:
                text_content = content.decode("utf-8", errors="ignore")
        else:
            text_content = content.decode("utf-8", errors="ignore")

        if not text_content.strip():
            raise HTTPException(400, "Could not extract text from document")

        analysis = await analyze_document(text_content)
        file_id = "doc_" + secrets.token_hex(6)
        doc = {
            "file_id": file_id,
            "filename": file.filename,
            "owner_id": current_user["user_id"],
            "content": text_content[:50000],
            "analysis": analysis,
            "size": len(content),
            "created_at": datetime.utcnow().isoformat(),
        }
        await db.documents.insert_one(doc)
        doc.pop("_id", None)
        doc.pop("content", None)
        return {"file_id": file_id, "filename": file.filename, "analysis": analysis}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(502, f"Document processing error: {str(e)}")

@router.post("/documents/query")
async def query_document(req: DocumentQueryRequest, current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    doc = await db.documents.find_one({"file_id": req.file_id, "owner_id": current_user["user_id"]})
    if not doc:
        raise HTTPException(404, "Document not found")
    try:
        answer = await analyze_document(doc["content"], req.question)
        return {"answer": answer, "question": req.question}
    except Exception as e:
        raise HTTPException(502, f"AI query error: {str(e)}")

@router.post("/documents/report")
async def generate_report(req: DocumentReportRequest, current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    if req.report_type not in REPORT_PROMPTS:
        raise HTTPException(400, f"Invalid report type. Choose from: {list(REPORT_PROMPTS.keys())}")
    doc = await db.documents.find_one({"file_id": req.file_id, "owner_id": current_user["user_id"]})
    if not doc:
        raise HTTPException(404, "Document not found")
    try:
        report = await generate_document_report(doc["content"], req.report_type)
        return {"report_type": req.report_type, "report": report, "filename": doc["filename"]}
    except Exception as e:
        raise HTTPException(502, f"Report generation error: {str(e)}")

@router.get("/documents/report-types")
async def list_report_types():
    return {"report_types": list(REPORT_PROMPTS.keys())}

@router.get("/documents")
async def list_documents(current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    cursor = db.documents.find({"owner_id": current_user["user_id"]}, {"_id": 0, "content": 0}).sort("created_at", -1)
    docs = await cursor.to_list(100)
    return {"documents": docs}

@router.delete("/documents/{file_id}")
async def delete_document(file_id: str, current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    await db.documents.delete_one({"file_id": file_id, "owner_id": current_user["user_id"]})
    return {"ok": True}

# ── Excel NLQ ─────────────────────────────────────────────────────────────────

@router.post("/excel/upload")
async def upload_excel(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    try:
        import pandas as pd
        content = await file.read()
        ext = (file.filename or "").split(".")[-1].lower()

        if ext == "csv":
            df = pd.read_csv(io.BytesIO(content))
        else:
            df = pd.read_excel(io.BytesIO(content))

        sheet_count = 1
        # For multi-sheet xlsx, count sheets
        if ext in ("xlsx", "xls"):
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True, read_only=True)
            sheet_count = len(wb.sheetnames)
            wb.close()

        columns = list(df.columns)
        numeric_cols = list(df.select_dtypes(include=["number"]).columns)
        row_count = len(df)

        # Summary text for NLQ
        summary_lines = [
            f"Columns ({len(columns)}): {', '.join(str(c) for c in columns)}",
            f"Rows: {row_count}",
            f"Numeric columns: {', '.join(str(c) for c in numeric_cols)}",
            "",
            "Sample (first 5 rows):",
        ]
        for _, row in df.head(5).iterrows():
            summary_lines.append("  " + " | ".join(str(v) for v in row))

        if numeric_cols:
            summary_lines.append("\nStatistics:")
            summary_lines.append(df[numeric_cols].describe().to_string())

        data_summary = "\n".join(summary_lines)

        # Store compact data for analysis (up to 1000 rows)
        data_records = df.head(1000).to_dict(orient="records")

        file_id = "xl_" + secrets.token_hex(6)
        await db.excel_files.insert_one({
            "file_id": file_id,
            "filename": file.filename,
            "owner_id": current_user["user_id"],
            "data_summary": data_summary,
            "data_records": data_records,
            "columns": columns,
            "numeric_columns": numeric_cols,
            "row_count": row_count,
            "sheet_count": sheet_count,
            "created_at": datetime.utcnow().isoformat(),
        })
        return {
            "file_id": file_id,
            "filename": file.filename,
            "columns": columns,
            "row_count": row_count,
            "sheet_count": sheet_count,
            "summary": data_summary[:500],
        }
    except Exception as e:
        raise HTTPException(502, f"Excel processing error: {str(e)}")

@router.post("/excel/query")
async def query_excel(req: ExcelQueryRequest, current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    xfile = await db.excel_files.find_one({"file_id": req.file_id, "owner_id": current_user["user_id"]})
    if not xfile:
        raise HTTPException(404, "File not found")
    try:
        answer = await excel_nlq(xfile["data_summary"], req.question)
        return {"answer": answer, "question": req.question}
    except Exception as e:
        raise HTTPException(502, f"AI query error: {str(e)}")

@router.post("/excel/analyze")
async def analyze_excel(req: ExcelQueryRequest, current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    """Full statistical analysis with AI insights"""
    xfile = await db.excel_files.find_one({"file_id": req.file_id, "owner_id": current_user["user_id"]})
    if not xfile:
        raise HTTPException(404, "File not found")
    try:
        import pandas as pd
        df = pd.DataFrame(xfile.get("data_records", []))
        stats_context = xfile["data_summary"]
        if not df.empty:
            numeric_cols = xfile.get("numeric_columns", [])
            if numeric_cols:
                stats_context += f"\n\nFull Statistics:\n{df[numeric_cols].describe().to_string()}"
            stats_context += f"\n\nColumn types:\n{df.dtypes.to_string()}"
        insights = await analyze_excel_data(stats_context)
        return {"insights": insights, "file_id": req.file_id}
    except Exception as e:
        raise HTTPException(502, f"Analysis error: {str(e)}")

@router.post("/excel/errors")
async def detect_errors(req: ExcelQueryRequest, current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    """Detect data quality issues and anomalies"""
    xfile = await db.excel_files.find_one({"file_id": req.file_id, "owner_id": current_user["user_id"]})
    if not xfile:
        raise HTTPException(404, "File not found")
    try:
        import pandas as pd
        df = pd.DataFrame(xfile.get("data_records", []))
        pre_detected = []

        if not df.empty:
            for col in df.columns:
                null_count = int(df[col].isnull().sum())
                if null_count > 0:
                    pre_detected.append({"column": col, "issue": "null_values", "count": null_count})

                if pd.api.types.is_numeric_dtype(df[col]) and len(df) > 3:
                    q1, q3 = df[col].quantile(0.25), df[col].quantile(0.75)
                    iqr = q3 - q1
                    outliers = int(((df[col] < q1 - 1.5 * iqr) | (df[col] > q3 + 1.5 * iqr)).sum())
                    if outliers > 0:
                        pre_detected.append({"column": col, "issue": "outliers", "count": outliers})

        data_preview = df.head(20).to_string() if not df.empty else "No data"
        ai_analysis = await detect_data_errors(data_preview, pre_detected)
        return {"detected_issues": pre_detected, "ai_analysis": ai_analysis, "total_issues": len(pre_detected)}
    except Exception as e:
        raise HTTPException(502, f"Error detection failed: {str(e)}")

@router.post("/excel/formula")
async def get_formula(req: FormulaRequest, current_user: dict = Depends(get_current_user)):
    """Generate Excel formula from natural language description"""
    try:
        result = await generate_formula(req.description, req.context)
        return {"formula_response": result}
    except Exception as e:
        raise HTTPException(502, f"Formula generation error: {str(e)}")

@router.get("/excel/files")
async def list_excel_files(current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    cursor = db.excel_files.find(
        {"owner_id": current_user["user_id"]},
        {"_id": 0, "data_summary": 0, "data_records": 0}
    ).sort("created_at", -1)
    files = await cursor.to_list(50)
    return {"files": files}

# ── Transcribe ────────────────────────────────────────────────────────────────

@router.post("/transcribe")
async def transcribe_audio(
    file: UploadFile = File(...),
    language: str = Form(default="en"),
    current_user: dict = Depends(get_current_user)
):
    try:
        content = await file.read()
        if not content:
            raise HTTPException(400, "Empty audio file")
        audio_url = await upload_audio(content)
        result = await transcribe_url(audio_url, language_code=language)
        if result.get("status") == "error":
            raise HTTPException(502, f"Transcription failed: {result.get('error', 'Unknown error')}")
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(502, f"Transcription service error: {str(e)}")

# ── Smart Files (folders + GridFS + AI chat) ──────────────────────────────────

def _get_gridfs(db):
    import motor.motor_asyncio
    return motor.motor_asyncio.AsyncIOMotorGridFSBucket(db)

async def _extract_text(content: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("txt", "md"):
        return content.decode("utf-8", errors="ignore")
    if ext == "pdf":
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""
    if ext in ("docx", "doc"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    if ext in ("xlsx", "xls", "csv"):
        try:
            import pandas as pd
            df = pd.read_csv(io.BytesIO(content)) if ext == "csv" else pd.read_excel(io.BytesIO(content))
            return df.to_string()
        except Exception:
            return ""
    if ext == "json":
        try:
            return json.dumps(json.loads(content.decode("utf-8")), indent=2)
        except Exception:
            return ""
    return ""

@router.post("/smart-files/folders")
async def create_folder(req: CreateFolderRequest, current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    folder = {
        "id": "folder_" + secrets.token_hex(6),
        "owner_id": current_user["user_id"],
        "name": req.name,
        "parent_id": req.parent_id,
        "type": "folder",
        "created_at": datetime.utcnow().isoformat(),
    }
    await db.smart_files.insert_one(folder)
    folder.pop("_id", None)
    return folder

@router.get("/smart-files/contents")
async def get_contents(
    folder_id: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    query = {"owner_id": current_user["user_id"], "parent_id": folder_id}
    items = await db.smart_files.find(query, {"_id": 0, "searchable_content": 0}).to_list(200)
    folders = [i for i in items if i.get("type") == "folder"]
    files = [i for i in items if i.get("type") == "file"]
    # Breadcrumb path
    path = []
    if folder_id:
        cur = folder_id
        for _ in range(10):
            f = await db.smart_files.find_one({"id": cur, "owner_id": current_user["user_id"]}, {"_id": 0})
            if not f:
                break
            path.insert(0, {"id": f["id"], "name": f["name"]})
            cur = f.get("parent_id")
            if not cur:
                break
    return {"folder_id": folder_id, "path": path, "folders": folders, "files": files}

@router.post("/smart-files/upload")
async def upload_smart_file(
    file: UploadFile = File(...),
    folder_id: Optional[str] = Form(default=None),
    description: str = Form(default=""),
    current_user: dict = Depends(get_current_user),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    try:
        content = await file.read()
        fs = _get_gridfs(db)
        from bson import ObjectId
        gridfs_id = await fs.upload_from_stream(
            file.filename,
            io.BytesIO(content),
            metadata={"owner_id": current_user["user_id"], "folder_id": folder_id}
        )
        text = await _extract_text(content, file.filename or "")
        file_meta = {
            "id": "sf_" + secrets.token_hex(6),
            "owner_id": current_user["user_id"],
            "gridfs_id": str(gridfs_id),
            "filename": file.filename,
            "file_type": (file.filename or "").rsplit(".", 1)[-1].lower(),
            "parent_id": folder_id,
            "description": description,
            "searchable_content": text[:50000],
            "size": len(content),
            "type": "file",
            "learned": bool(text),
            "created_at": datetime.utcnow().isoformat(),
        }
        await db.smart_files.insert_one(file_meta)
        file_meta.pop("_id", None)
        file_meta.pop("searchable_content", None)
        return file_meta
    except Exception as e:
        raise HTTPException(502, f"Upload error: {str(e)}")

@router.get("/smart-files/download/{file_id}")
async def download_smart_file(file_id: str, current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    meta = await db.smart_files.find_one({"id": file_id, "owner_id": current_user["user_id"], "type": "file"})
    if not meta:
        raise HTTPException(404, "File not found")
    try:
        from bson import ObjectId
        fs = _get_gridfs(db)
        stream = await fs.open_download_stream(ObjectId(meta["gridfs_id"]))
        data = await stream.read()
        return StreamingResponse(io.BytesIO(data), media_type="application/octet-stream",
            headers={"Content-Disposition": f'attachment; filename="{meta["filename"]}"'})
    except Exception as e:
        raise HTTPException(502, f"Download error: {str(e)}")

@router.delete("/smart-files/{item_id}")
async def delete_smart_item(item_id: str, current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    item = await db.smart_files.find_one({"id": item_id, "owner_id": current_user["user_id"]})
    if not item:
        raise HTTPException(404, "Item not found")
    try:
        from bson import ObjectId
        if item["type"] == "file" and item.get("gridfs_id"):
            fs = _get_gridfs(db)
            await fs.delete(ObjectId(item["gridfs_id"]))
        elif item["type"] == "folder":
            # Delete all files inside recursively
            async def delete_children(fid):
                children = await db.smart_files.find({"owner_id": current_user["user_id"], "parent_id": fid}).to_list(1000)
                for c in children:
                    if c["type"] == "folder":
                        await delete_children(c["id"])
                    elif c.get("gridfs_id"):
                        fs = _get_gridfs(db)
                        await fs.delete(ObjectId(c["gridfs_id"]))
                await db.smart_files.delete_many({"owner_id": current_user["user_id"], "parent_id": fid})
            await delete_children(item_id)
        await db.smart_files.delete_one({"id": item_id})
        return {"ok": True}
    except Exception as e:
        raise HTTPException(502, f"Delete error: {str(e)}")

@router.post("/smart-files/chat")
async def smart_files_ai_chat(req: SmartFilesChatRequest, current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    learned = await db.smart_files.find(
        {"owner_id": current_user["user_id"], "type": "file", "learned": True},
        {"filename": 1, "searchable_content": 1, "_id": 0}
    ).to_list(30)
    if not learned:
        return {"response": "No files indexed yet. Upload files and they'll be automatically indexed for AI chat.", "files_used": 0}
    try:
        contexts = [{"filename": f["filename"], "content": f.get("searchable_content", "")} for f in learned if f.get("searchable_content")]
        response = await smart_files_chat(req.message, contexts)
        return {"response": response, "files_used": len(contexts)}
    except Exception as e:
        raise HTTPException(502, f"AI chat error: {str(e)}")

# ── Smart Files — combined listing ────────────────────────────────────────────

@router.get("/files/list")
async def list_files(current_user: dict = Depends(get_current_user), db: AsyncIOMotorDatabase = Depends(get_db)):
    uid = current_user["user_id"]
    docs_cursor = db.documents.find({"owner_id": uid}, {"_id": 0, "content": 0}).sort("created_at", -1).limit(20)
    xl_cursor = db.excel_files.find({"owner_id": uid}, {"_id": 0, "data_summary": 0, "data_records": 0}).sort("created_at", -1).limit(20)
    docs = await docs_cursor.to_list(20)
    xls = await xl_cursor.to_list(20)
    for d in docs:
        d["type"] = "document"
    for x in xls:
        x["type"] = "excel"
    all_files = sorted(docs + xls, key=lambda f: f.get("created_at", ""), reverse=True)
    return {"files": all_files}
